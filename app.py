import os
import uuid
from docx import Document
from flask import Flask, request, jsonify, Response
from transformers import pipeline
import yt_dlp
import requests

# Initialize Flask app
app = Flask(__name__)

# Store the transcription dynamically
transcription = ""

# Load the Hugging Face model pipeline locally
qa_pipeline = pipeline("text-generation", model="gpt2")

# Route for solving questions (generating answers using GPT-2 model)
@app.route("/solve", methods=["POST"])
def solve():
    data = request.get_json()
    description = data.get("description", "")  # Extract description from request

    if not description:
        return jsonify({"error": "Description is missing"}), 400

    # Use the model to generate a solution
    result = qa_pipeline(description, max_length=100, num_return_sequences=1)

    return jsonify({"answer": result[0]["generated_text"]})

# Route for YouTube transcription (generating transcript and saving it as Word document)
@app.route("/transcribe", methods=["POST"])
def transcribe():
    global transcription
    data = request.get_json()
    youtube_url = data.get("url", "")
    course_id = data.get("course_id", "")
    title = data.get("title", "Lesson from YouTube")

    if not youtube_url:
        return jsonify({"error": "YouTube URL is missing"}), 400

    try:
        # Create a unique filename for the audio
        filename = f"{uuid.uuid4()}.mp3"

        # Download audio from YouTube
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': filename,
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            'quiet': False  # Set to False to get more logs
        }

        print("Starting to download the audio...")
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([youtube_url])

        # Simulate the result (replace this with actual transcription logic)
        print("Transcribing the audio...")
        result = {"text": "This is a simulated transcript of the video."}
        transcription = result['text']

        # Create a Word document with the transcript
        doc = Document()
        doc.add_heading('Transcription of Video', 0)
        doc.add_paragraph(transcription)

        # Save the document in the static folder with a unique filename
        doc_filename = f"{uuid.uuid4()}.docx"
        doc.save(os.path.join('static', 'study_materials', doc_filename))

        # Clean up the downloaded audio file
        os.remove(filename)

        # Send the generated Word file to Express server for storing metadata
        express_url = "http://localhost:5000/admin/add-content"  # URL to your Express backend
        lesson_data = {
            "courseId": course_id,
            "title": title,
            "downloadableMaterial": doc_filename  # Send the Word file's name as a reference
        }

        response = requests.post(express_url, json=lesson_data)

        if response.status_code == 200:
            return jsonify({"transcript": transcription, "message": "Lesson added successfully", "file": doc_filename})
        else:
            print(f"Failed to add lesson to server: {response.text}")
            return jsonify({"error": "Failed to add lesson to the server"}), 500

    except Exception as e:
        print(f"Error during transcription: {str(e)}")
        return jsonify({"error": "Failed to fetch transcription", "details": str(e)}), 500

# Route for serving the transcription as a download (text file)
@app.route("/download/")
def download():
    if not transcription:
        return jsonify({"error": "No transcription available"}), 400

    # Create the response with the transcription as a text file
    response = Response(transcription, mimetype='text/plain')
    response.headers["Content-Disposition"] = "attachment; filename=transcription.txt"
    return response

if __name__ == "__main__":
    app.run(debug=True, port=5001)
